#!/usr/bin/env python3
import csv
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / 'paper' / 'SAB_HET_2.0_edited.docx'
TABLE2 = ROOT / 'results' / 'paper' / 'final' / 'tables' / 'Table2_ordinal_PO_vs_nonPO_power_bias.tsv'


def read_tsv(path: Path):
    with path.open('r', encoding='utf-8', newline='') as f:
        r = csv.DictReader(f, delimiter='\t')
        rows = list(r)
    return r.fieldnames, rows


def insert_table_after_paragraph(doc: Document, paragraph, rows, headers):
    # Create table at end, then move XML after paragraph
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # header
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    # body
    for i, row in enumerate(rows, start=1):
        for j, h in enumerate(headers):
            table.cell(i, j).text = str(row.get(h, ''))

    tbl = table._element
    p = paragraph._element
    p.addnext(tbl)
    return table


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    headers, rows = read_tsv(TABLE2)

    # backup
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup = DOC_PATH.with_name(f'{DOC_PATH.stem}.backup_{stamp}.docx')
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))

    # If Table 2 already inserted, skip
    for p in doc.paragraphs:
        if p.text.strip().startswith('Table 2:'):
            print('Table 2 caption already present, no insertion performed.')
            doc.save(str(DOC_PATH))
            print(f'Backup: {backup}')
            return

    # Find Figure 4 caption paragraph
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('Figure 4:'):
            target = p
            break

    if target is None:
        raise RuntimeError('Could not find Figure 4 caption paragraph to insert after.')

    caption = (
        'Table 2: Ordinal PO vs non-PO performance at N=3,000, showing power, '
        'Type S, Type M, proportional bias, and RMSE for binary (death) and ordinal models '
        'under 70% and 100% accuracy.'
    )

    # Insert caption paragraph after target
    # add new paragraph at end then move
    new_p = doc.add_paragraph(caption)
    new_p._element.getparent().remove(new_p._element)
    target._element.addnext(new_p._element)

    # Insert table after caption
    insert_table_after_paragraph(doc, new_p, rows, headers)

    doc.save(str(DOC_PATH))
    print(f'Inserted Table 2. Updated: {DOC_PATH}')
    print(f'Backup: {backup}')


if __name__ == '__main__':
    main()
