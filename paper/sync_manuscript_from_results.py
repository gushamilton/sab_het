#!/usr/bin/env python3
"""Sync key manuscript placeholders with current simulation outputs.

Updates paper/SAB_HET_2.0_edited.docx in place (with timestamped backup).
"""

import csv
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "paper" / "SAB_HET_2.0_edited.docx"
MAIN_METRICS = ROOT / "results" / "main" / "data" / "main_binary_metrics.tsv"
TABLE2_PATH = ROOT / "results" / "paper" / "final" / "tables" / "Table2_ordinal_PO_vs_nonPO_power_bias.tsv"


def read_tsv(path: Path):
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f, delimiter="\t"))


def fmt3(x: float) -> str:
    return f"{x:.3f}"


def insert_paragraph_before(target_paragraph, text: str):
    new_p = target_paragraph._parent.add_paragraph(text)
    target_paragraph._p.addprevious(new_p._p)
    return new_p


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"Missing manuscript doc: {DOC_PATH}")

    main_rows = read_tsv(MAIN_METRICS)
    table2_rows = read_tsv(TABLE2_PATH)

    # Lookup: power by sample size/group for accuracy=1, alpha=0.05.
    power_lookup = {}
    for row in main_rows:
        if float(row["alpha"]) != 0.05:
            continue
        if float(row["accuracy"]) != 1.0:
            continue
        grp = row["group"]
        n = int(float(row["sample_size"]))
        p = row["power"]
        if grp != "A" and p not in ("", "NA", None):
            power_lookup[(n, grp)] = float(p)

    # Key numeric summaries for text insertion.
    getp = lambda n, g: power_lookup[(n, g)]
    p_b_500 = getp(500, "B")
    p_b_5000 = getp(5000, "B")
    p_b_10000 = getp(10000, "B")
    p_b_20000 = getp(20000, "B")
    p_c_20000 = getp(20000, "C")
    p_d_20000 = getp(20000, "D")
    p_e_20000 = getp(20000, "E")

    type1_a_20000 = [
        float(r["type1"])
        for r in main_rows
        if int(float(r["sample_size"])) == 20000
        and r["group"] == "A"
        and float(r["alpha"]) == 0.05
        and r["type1"] not in ("", "NA", None)
    ]
    type1_min = min(type1_a_20000)
    type1_max = max(type1_a_20000)

    # Table 2 (PO/nonPO at N=3000, acc=1) summary.
    t2 = [
        r
        for r in table2_rows
        if int(float(r["sample_size"])) == 3000
        and float(r["accuracy"]) == 1.0
        and r["group"] in {"B", "C", "D", "E"}
    ]

    def pick(dgm: str, group: str, model: str):
        return next(
            r for r in t2 if r["dgm"] == dgm and r["group"] == group and r["model_type"] == model
        )

    po_b_bin = float(pick("PO", "B", "Binary (death)")["power"])
    po_b_ord = float(pick("PO", "B", "Ordinal (polr)")["power"])
    po_c_bin = float(pick("PO", "C", "Binary (death)")["power"])
    po_c_ord = float(pick("PO", "C", "Ordinal (polr)")["power"])
    po_d_bin = float(pick("PO", "D", "Binary (death)")["power"])
    po_d_ord = float(pick("PO", "D", "Ordinal (polr)")["power"])
    po_e_bin = float(pick("PO", "E", "Binary (death)")["power"])
    po_e_ord = float(pick("PO", "E", "Ordinal (polr)")["power"])

    np_b_bin = float(pick("nonPO", "B", "Binary (death)")["power"])
    np_b_ord = float(pick("nonPO", "B", "Ordinal (polr)")["power"])
    np_c_bin = float(pick("nonPO", "C", "Binary (death)")["power"])
    np_c_ord = float(pick("nonPO", "C", "Ordinal (polr)")["power"])
    np_d_bin = float(pick("nonPO", "D", "Binary (death)")["power"])
    np_d_ord = float(pick("nonPO", "D", "Ordinal (polr)")["power"])
    np_e_bin = float(pick("nonPO", "E", "Binary (death)")["power"])
    np_e_ord = float(pick("nonPO", "E", "Ordinal (polr)")["power"])

    # Backup then edit.
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC_PATH.with_name(f"{DOC_PATH.stem}.backup_{stamp}.docx")
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))

    add_data = (
        f"At 100% classification accuracy, power for subgroup B rose from {fmt3(p_b_500)} (n=500) "
        f"to {fmt3(p_b_5000)} (n=5,000), {fmt3(p_b_10000)} (n=10,000), and {fmt3(p_b_20000)} (n=20,000); "
        f"however, at n=20,000 power remained below 50% for C ({fmt3(p_c_20000)}), "
        f"D ({fmt3(p_d_20000)}), and E ({fmt3(p_e_20000)})."
    )

    add_type1 = (
        f"Type I error in the null subgroup (A) remained close to nominal in the main simulation "
        f"(n=20,000), ranging from {100*type1_min:.1f}% at 100% accuracy to {100*type1_max:.1f}% at 70% accuracy."
    )

    add_ord = (
        f"At n=3,000 and 100% accuracy, under proportional odds the ordinal model was more powerful than binary in all non-null subgroups: "
        f"B ({fmt3(po_b_ord)} vs {fmt3(po_b_bin)}), C ({fmt3(po_c_ord)} vs {fmt3(po_c_bin)}), "
        f"D ({fmt3(po_d_ord)} vs {fmt3(po_d_bin)}), and E ({fmt3(po_e_ord)} vs {fmt3(po_e_bin)}). "
        f"Under non-proportional death-only effects, binary death analysis was more powerful for B ({fmt3(np_b_bin)} vs {fmt3(np_b_ord)}), "
        f"C ({fmt3(np_c_bin)} vs {fmt3(np_c_ord)}), and D ({fmt3(np_d_bin)} vs {fmt3(np_d_ord)}), "
        f"whereas ordinal was only slightly higher in E ({fmt3(np_e_ord)} vs {fmt3(np_e_bin)})."
    )

    for p in doc.paragraphs:
        txt = p.text
        if "hypothetical, fixed, large trial of 10,000 participants" in txt:
            p.text = txt.replace("hypothetical, fixed, large trial of 10,000 participants", "hypothetical, fixed, large trial of 20,000 participants")
            txt = p.text
        if "As shown in Figure 3" in txt and "classification accuracy" in txt:
            p.text = txt.replace("As shown in Figure 3", "As shown in Figure 1")
            txt = p.text
        if "Power estimated from 500 replicates" in txt:
            p.text = txt.replace("Power estimated from 500 replicates", "Power estimated from 2,000 replicates")
            txt = p.text
        if "[ADD DATA]" in txt:
            p.text = txt.replace("[ADD DATA]", add_data)
            txt = p.text
        if txt.strip() == "ADD TYPE1 ERROR":
            p.text = add_type1
            txt = p.text
        if txt.strip() == "ADD TABLE OF BIAS/power/":
            p.text = add_ord
            txt = p.text
        if "We compared ordinal versus binary analyses for power (n = 2,500)" in txt:
            p.text = txt.replace("power (n = 2,500)", "power (n = 3,000)")
            txt = p.text
        if txt.startswith("Finally, we tested whether ordinal outcomes could rescue power lost by binary analyses in this setting."):
            p.text = (
                "Finally, we tested whether ordinal outcomes could rescue power lost by binary analyses in this setting. "
                "We simulated a 6-point ordinal outcome using the same subphenotype setup as in ARREST under two scenarios: "
                "(1) a proportional-odds shift across the outcome scale, and (2) a non-proportional death-only effect "
                "(treatment changes death odds, with non-death categories rescaled proportionally). "
                "The corresponding category shifts for a representative subgroup are shown in Figure 4."
            )
            txt = p.text
        if "When the ordinal outcome was appropriate (e.g. proportional odds), power was dramatically increased using an ordinal outcome, while type 1 error control was maintained (Figure 4A,4B)." in txt:
            p.text = txt.replace("Figure 4A,4B", "Figure 5A,5B")
            txt = p.text
        if txt.startswith("Figure 4:"):
            p.text = (
                "Figure 4: Ordinal outcome category shifts under proportional-odds (A) and non-proportional "
                "(death-only) effects (B) for a representative subgroup. Stacked bars show the control and treatment "
                "distributions across the six ordered categories (death, ICU, hospital, complications, discharged, recovered). "
                "Dashed connectors link matching categories between arms to show how probability mass moves across the outcome scale."
            )

    # Insert Figure 5 caption before Table 2 so the manuscript reads Figure 4 -> Figure 5 -> Table 2.
    if not any((p.text or "").strip().startswith("Figure 5:") for p in doc.paragraphs):
        fig5_caption = (
            "Figure 5: Binary versus ordinal performance under proportional-odds and non-proportional "
            "(death-only) ordinal data-generating mechanisms, stratified by classification accuracy. "
            "Panels A and C show replicate scaled error at n=20,000; panels B and D show power comparison at n=3,000."
        )
        table2_para = next(
            (p for p in doc.paragraphs if (p.text or "").strip().startswith("Table 2: Ordinal PO vs non-PO performance")),
            None,
        )
        if table2_para is None:
            raise RuntimeError("Could not locate Table 2 caption to insert Figure 5 before.")
        insert_paragraph_before(table2_para, fig5_caption)

    # Fill Table 3 power grid in manuscript (second table).
    t = doc.tables[1]
    row_groups = ["A", "B", "C", "D", "E"]
    for ridx in range(2, len(t.rows)):
        n_txt = t.cell(ridx, 0).text.strip()
        if not n_txt.isdigit():
            continue
        n = int(n_txt)
        for cidx, g in enumerate(row_groups, start=1):
            if g == "A":
                t.cell(ridx, cidx).text = "-"
            else:
                val = power_lookup.get((n, g))
                t.cell(ridx, cidx).text = fmt3(val) if val is not None else "-"

    doc.save(str(DOC_PATH))
    print(f"Updated: {DOC_PATH}")
    print(f"Backup:  {backup}")


if __name__ == "__main__":
    main()
