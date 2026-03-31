#!/usr/bin/env python3
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[2]
FIG_DIR = ROOT / 'paper' / 'fig_tmp'


def resolve_doc_path(root: Path) -> Path:
    paper_dir = root / 'paper'
    env_path = Path(__import__('os').environ['MANUSCRIPT_DOC']) if 'MANUSCRIPT_DOC' in __import__('os').environ else None
    if env_path is not None:
        return env_path

    preferred = paper_dir / 'SAB_HET_2.1_JU.docx'
    if preferred.exists():
        return preferred

    preferred = paper_dir / 'SAB_HET_2.0_edited.docx'
    if preferred.exists():
        return preferred

    backups = sorted(
        paper_dir.glob('SAB_HET_2.0_edited.backup_*.docx'),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    if backups:
        return backups[0]

    fallback = paper_dir / 'SAB_HET_2.0.docx'
    if fallback.exists():
        return fallback

    raise FileNotFoundError(f'No manuscript doc found in {paper_dir}')


DOC_PATH = resolve_doc_path(ROOT)

FIG_MAP = {
    'Figure 1:': FIG_DIR / 'Figure1.png',
    'Figure 2:': FIG_DIR / 'Figure2.png',
    'Figure 3:': FIG_DIR / 'Figure3.png',
    'Figure 4:': FIG_DIR / 'Figure4.png',
    'Figure 5:': FIG_DIR / 'Figure5.png',
}


def remove_all_inline_drawings(doc: Document):
    # Remove all runs that contain drawing nodes.
    for p in doc.paragraphs:
        for r in list(p.runs):
            if r._r.xpath('.//w:drawing'):
                r._r.getparent().remove(r._r)


def insert_picture_after_paragraph(doc: Document, paragraph, image_path: Path, width_in=6.5):
    new_p = doc.add_paragraph()
    run = new_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    new_p._element.getparent().remove(new_p._element)
    paragraph._element.addnext(new_p._element)


def main():
    for k, path in FIG_MAP.items():
        if not path.exists():
            raise FileNotFoundError(f'Missing image for {k}: {path}')

    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup = DOC_PATH.with_name(f'{DOC_PATH.stem}.backup_{stamp}.docx')
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))

    remove_all_inline_drawings(doc)

    inserted = 0
    for p in doc.paragraphs:
        txt = (p.text or '').strip()
        for prefix, img in FIG_MAP.items():
            if txt.startswith(prefix):
                insert_picture_after_paragraph(doc, p, img)
                inserted += 1
                break

    doc.save(str(DOC_PATH))
    print(f'Updated: {DOC_PATH}')
    print(f'Backup:  {backup}')
    print(f'Inserted figure images: {inserted}')


if __name__ == '__main__':
    main()
