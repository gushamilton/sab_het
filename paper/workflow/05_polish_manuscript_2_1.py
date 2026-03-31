#!/usr/bin/env python3
import os
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]


def resolve_doc_path() -> Path:
    env_doc = os.environ.get("MANUSCRIPT_DOC")
    if env_doc:
        return Path(env_doc)
    return ROOT / "paper" / "SAB_HET_2.1_JU.docx"


DOC = resolve_doc_path()


def set_run_font(run, font_name: str = "Arial") -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font_name)


def apply_arial(doc: Document) -> None:
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        except Exception:
            pass

    for para in doc.paragraphs:
        for run in para.runs:
            set_run_font(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run)


def main() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC.with_name(f"{DOC.stem}.backup_{stamp}.docx")
    shutil.copy2(DOC, backup)

    replacements = {
        "Staphylococcus aureus bacteraemia (SAB) is clinically heterogeneous. Recently identified, reproducible subphenotypes, that is, recurrent clinical patient groups with distinct characteristics and outcomes, suggest potential for heterogeneous treatment effects (HTE), but the impact of patient misclassification on detecting these effects is unclear, and strategies to improve detection of HTE are not yet identified.": (
            "Staphylococcus aureus bacteraemia (SAB) is clinically heterogeneous. Recently identified and reproducible subphenotypes, meaning recurrent clinical patient groups with distinct characteristics and outcomes, suggest the potential for heterogeneous treatment effects (HTE). However, the impact of misclassifying patients into these groups is unclear, and practical strategies to improve HTE detection remain uncertain."
        ),
        "Even with perfect classification, post-hoc detection of heterogeneous treatment effects remained highly conditional on subphenotype prevalence, baseline mortality, and effect size. Subphenotype B was readily detectable at moderate sample sizes, whereas C and D remained challenging and E remained difficult under binary mortality analyses alone. Decreasing classification accuracy reduced power, increased type I error, and introduced bias. Enrichment appeared realistic for subphenotype B but remained unattractive for most other subphenotypes. Ordinal outcomes substantially improved performance when they matched the treatment-effect structure, but this depended on both the assumed mechanism and the number of ordinal levels used.": (
            "Even with perfect classification, post-hoc detection of heterogeneous treatment effects remained highly conditional on subgroup prevalence, baseline mortality, and effect size. One subgroup, subphenotype B, was readily detectable at moderate sample sizes, whereas C and D remained challenging and E remained difficult under binary mortality analyses alone. Decreasing classification accuracy reduced power, increased type I error, and introduced bias. Enrichment appeared realistic for subphenotype B but remained unattractive for most other subphenotypes. Ordinal outcomes substantially improved performance when they matched the treatment-effect structure, but this depended on both the assumed mechanism and the number of ordinal levels used."
        ),
        "We performed a simulation study using recent data from randomised trials and observational studies in SAB. We assessed the impact classification accuracy (70%-100%) on i) power, ii) type 1 error, and iii) bias in post-hoc analyses of HTE. We modelled two strategies to try and maximise power and reduce bias; first by performing enrichment trials (e.g. only randomising those who are predicted to be in a subphenotype), and secondly, by using ordinal outcomes rather than binary outcomes.": (
            "We performed a simulation study using recent data from randomised trials and observational studies in SAB. We assessed the impact of classification accuracy (70%-100%) on i) power, ii) type 1 error, and iii) bias in post-hoc analyses of HTE. We then evaluated two strategies to improve performance: enrichment designs, in which only patients predicted to belong to a target subphenotype are randomised, and the use of ordinal rather than binary outcomes."
        ),
        "Recent efforts have focused on identifying clinically relevant subphenotypes within SAB to enable better patient stratification for research and potentially personalized treatment. Using latent class analysis on routinely available data from observational and trial cohorts from the UK, Spain, the Netherlands, and USA five distinct and reproducible clinical subphenotypes (A-E) have been identified (and replicated).8–10": (
            "Recent efforts have focused on identifying clinically relevant subphenotypes within SAB to enable better patient stratification for research and potentially personalised treatment. Using latent class analysis on routinely available data from observational and trial cohorts from the UK, Spain, the Netherlands, and the USA, investigators identified five distinct and reproducible clinical subphenotypes (A-E).8–10"
        ),
        "The possibility of such HTE raises critical questions for future clinical trial design and the implementation of stratified medicine approaches. In principle, participants in an RCT (or in clinical practice) could be subphenotyped, then assigned a treatment likely to specifically benefit them. This approach may be more beneficial than applying therapies across unselected patients with SAB, especially when for many subphenotypes the biological plausibility of some therapies may be much less. For example, in nosocomial catheter associated bacteraemia, where the primary treatment is catheter removal, the benefit of adjunctive clindamycin may be much less than in patients with a much higher burden of infection such as patients in with higher infection severity or bioburden such as in endocarditis and native vertebral osteomyelitis.": (
            "The possibility of such HTE raises important questions for future trial design and for stratified medicine in SAB. In principle, participants in an RCT, or patients in clinical practice, could be subphenotyped and then offered treatments more likely to help them. This may be preferable to applying the same therapy across an unselected SAB population, particularly when the biological plausibility of benefit differs across subphenotypes. For example, in nosocomial catheter-associated bacteraemia, where catheter removal is the main intervention, the potential benefit of adjunctive antimicrobial strategies may be smaller than in patients with a higher burden of infection, such as those with endocarditis or native vertebral osteomyelitis."
        ),
        "In our main analysis we simulated a simple, two arm randomized trial with subgroup-specific treatment effects derived from an analysis of the ARREST trial to get base estimates of the sample size required to identify sub-phenotypes in SAB. We then imposed imperfect subphenotype classification, and estimated subgroup effects using observed (potentially misclassified) subphenotypes. Table 1 describes the main simulated estimation parameters; these were taken from the frequencies and baseline mortality of the sub-phenotypes derived from ARREST. As the effect sizes were likely inflated by Winner’s curse, these were shrunk by 50% to more conservative effects for the simulations.12": (
            "In our main analysis we simulated a simple two-arm randomised trial with subgroup-specific treatment effects derived from analyses of the ARREST trial, to obtain baseline estimates of the sample size required to detect subphenotype-specific effects in SAB. We then imposed imperfect subphenotype classification and estimated subgroup effects using the observed, potentially misclassified, labels. Table 1 summarises the main simulation parameters, including subphenotype frequencies and baseline mortality derived from ARREST. Because the subgroup effect sizes were likely inflated by winner's curse, we shrank them by 50% on the log-odds scale for the primary simulations.12"
        ),
        "Across cohorts where subphenotypes of SAB have been measured, subphenotype prevalence, and mortality within each subphenotype differs greatly. We tested how sensitive our simulations were to these changes in prevalence, while holding the simulated effect size at the same level.": (
            "Across cohorts in which SAB subphenotypes have been measured, both subphenotype prevalence and within-subphenotype mortality vary substantially. We therefore tested how sensitive the simulations were to these changes in prevalence and mortality while holding the subgroup treatment effect constant."
        ),
        "The statistical power for detecting subphenotype-specific treatment effects in a standard two-arm RCT, assuming perfect classification and the conservatively shrunk effects used here, was highly dependent on subphenotype prevalence, baseline event rate, and effect size. The results, summarized in Table 3 and Figure 1, show that subphenotype B was readily detectable at realistic sample sizes, reaching essentially complete power by n=5,000. In contrast, C and D remained only modestly powered even at larger sample sizes, and E remained difficult under binary mortality analyses despite becoming clearly more plausible than in earlier iterations of the simulation. As expected in a traditional two-arm RCT, Type I error was well calibrated and bias was minimal.": (
            "The statistical power to detect subphenotype-specific treatment effects in a standard two-arm RCT, assuming perfect classification and the conservatively shrunk effects used here, was highly dependent on subphenotype prevalence, baseline event rate, and effect size. The results, summarised in Table 3 and Figure 1, show that subphenotype B was readily detectable at realistic sample sizes, reaching essentially complete power by n=5,000. In contrast, C and D remained only modestly powered even at larger sample sizes, and E remained difficult under binary mortality analyses despite becoming clearly more plausible than in earlier iterations of the simulation. As expected in a traditional two-arm RCT, Type I error was well calibrated and bias was minimal."
        ),
        "Table 3: Statistical power to detect a subphenotype specific effect given a sample size n in each subphenotype. Power estimated from 2,000 replicates where simulated. Subphenotype A has no treatment effect so power is not the appropriate metric.": (
            "Table 3: Statistical power to detect a subphenotype-specific effect at each sample size n. Power was estimated from 2,000 simulated replicates. Subphenotype A has no treatment effect, so power is not the appropriate metric."
        ),
        "We then investigated the impact of imperfect classification accuracy on power, Type 1 error, bias and bias in a hypothetical, fixed, large trial of 20,000 participants. As shown in Figure 1, statistical power to detect true effects diminished substantially as classification accuracy decreased. For null subgroups like A, this led to Type 1 error, while for all subgroups effect estimates were biased. The degree of bias related to the amount of misclassification from other groups (e.g. the accuracy of the test), and on the make-up of other groups for sub-phenotypes with a non-null effect. This \"dilution\" effect occurs because misclassifying patients mixes individuals from subphenotypes with different true effects, pulling the observed estimate towards the population average. This effect was larger for less frequent subphenotypes with large effects: the powerful effect in subphenotype B (OR=4.3) was severely attenuated with even very small misclassification.": (
            "We then investigated the impact of imperfect classification accuracy on power, Type 1 error, and bias in a hypothetical large trial of 20,000 participants. As shown in Figure 1, statistical power to detect true effects diminished substantially as classification accuracy decreased. For null subgroups such as A, this increased Type 1 error, while for all subgroups effect estimates became biased. The degree of bias depended both on the amount of misclassification and on the composition of the groups contributing those misclassified patients. This dilution effect occurs because misclassification mixes patients from subphenotypes with different true effects, pulling the observed estimate towards the population average. The effect was particularly marked for less frequent subphenotypes with large effects: the strong effect in subphenotype B (OR=4.3 after shrinkage) was substantially attenuated by even small amounts of misclassification."
        ),
        "It is important to understand how conditional these estimates are on a specific cohort because sample size calculations are made before the trial starts, and it is not known how much subphenotype prevalence and mortality variation between cohorts might influence outcomes. We therefore reran analyses using recent data from numerous cohorts with differing subphenotype prevalences, but with perfect classification. These show that for most subphenotypes, the sample size required varied by setting, by as much as 5-fold (Figure 2).": (
            "It is important to understand how conditional these estimates are on the cohort under study, because sample-size calculations are made before a trial starts and the extent to which subphenotype prevalence and mortality vary across settings is uncertain. We therefore reran the analyses using data from several cohorts with differing subphenotype prevalences and mortality, assuming perfect classification. For most subphenotypes, the required sample size varied substantially by setting, in some cases by as much as five-fold (Figure 2)."
        ),
        "These findings identify that even establishing the a priori sample size calculation for detection of HTE is extremely challenging, as it is unlikely that the proposed trial setting will have a similar sub-phenotype distribution than historic data. Supporting this, the predicted total sample size requirement for the two Edinburgh cohorts was quite different, reflecting small but measurable differences in subphenotype prevalence and mortality.": (
            "These findings show that even establishing an a priori sample-size calculation for HTE detection is difficult, because the planned trial setting is unlikely to mirror the subphenotype distribution seen in historical data. Supporting this, the predicted total sample-size requirement differed meaningfully even between the two Edinburgh cohorts, reflecting relatively small but still important differences in subphenotype prevalence and mortality."
        ),
        "Above, we demonstrated the magnitude of the subphenotype-specific treatment effect, the frequency of the target subphenotype in the population, and the baseline event rate. We took this further in a hypothetical exercise of ‘screening’ patients with SAB with an imperfect test to randomise them prospectively, as might be considered for a new trial.": (
            "We then extended this by considering a hypothetical enrichment design in which patients with SAB are screened using an imperfect classifier and only those predicted to belong to a target subphenotype are randomised prospectively."
        ),
        "Ordinal outcomes can provide greater precision when they capture clinically meaningful gradations of severity. In this analysis we compared binary versus ordinal endpoints under the same subgroup-effect structure and misclassification settings to assess gains in precision and their relationship to bias. Specifically, we used a 6-point ordinal scale (death, ICU, hospital, complications, discharged, recovered). We evaluated two scenarios: one with a proportional-odds effect across the full ordinal scale, and one with a non-proportional death-only effect in which treatment changes death odds while non-death categories are rescaled proportionally.": (
            "Ordinal outcomes can provide greater precision when they capture clinically meaningful gradations of severity. In this analysis we compared binary and ordinal endpoints under the same subgroup-effect structure and misclassification settings to assess gains in precision and their relationship to bias. Specifically, we used a 6-point ordinal scale spanning death to recovery. We evaluated two scenarios: one with a proportional-odds effect across the full ordinal scale, and one with a non-proportional death-only effect in which treatment changes death odds while non-death categories are rescaled proportionally."
        ),
        "Figure 5: Binary versus ordinal performance under proportional-odds and non-proportional (death-only) ordinal data-generating mechanisms, stratified by classification accuracy. Panels A and C show replicate scaled error at n=20,000; panels B and D show power comparison at n=3,000.": (
            "Figure 5: Binary versus ordinal performance under proportional-odds and non-proportional (death-only) ordinal data-generating mechanisms, stratified by classification accuracy. Panels A and C show replicate scaled error at n=20,000; panels B and D show power comparisons at n=3,000."
        ),
        "These results show the strengths and tradeoffs of ordinal outcomes. When the ordinal structure reasonably reflects the treatment effect, they can be much more powerful and less biased than binary mortality analyses. When the effect is concentrated mainly at death, binary analysis remains preferable, but the degree of ordinal underperformance depends on how the ordinal scale is constructed. In particular, our 6-point main analysis performed slightly better under proportional-odds scenarios, whereas the 5-point sensitivity analysis was somewhat more robust under death-only non-proportional scenarios.": (
            "These results show the strengths and tradeoffs of ordinal outcomes. When the ordinal structure reasonably reflects the treatment effect, they can be much more powerful and less biased than binary mortality analyses. When the effect is concentrated mainly at death, binary analysis remains preferable, but the degree of ordinal underperformance depends on how the ordinal scale is constructed. In particular, our 6-point main analysis performed slightly better under proportional-odds scenarios, whereas the 5-point sensitivity analysis was somewhat more robust under death-only non-proportional scenarios. This tradeoff is shown directly in Supplementary Figure S1."
        ),
        "This simulation study demonstrates that detecting heterogeneous treatment effects in SAB subphenotypes is substantially more challenging than current literature might suggest. The feasibility of identifying differential treatment effects is dictated not simply by the magnitude of a putative effect, but by the interplay of subphenotype prevalence, baseline event rate, classification accuracy, and outcome choice.": (
            "This simulation study shows that detecting heterogeneous treatment effects in SAB subphenotypes is more challenging than the headline subgroup estimates in the current literature might suggest. Feasibility is determined not simply by the magnitude of a putative effect, but by the interplay of subphenotype prevalence, baseline event rate, classification accuracy, and outcome choice."
        ),
        "We do not think this complexity should discourage core outcome sets for SAB. Standardised outcomes remain essential for comparability and should be reported in all trials. However, our data suggests they should function as a minimum reporting standard rather than a constraint on what the primary endpoint should be. Trialists should be encouraged to select a primary outcome matched to the hypothesised mechanism, while reporting the core set for cross-trial comparison. This distinction between \"always report\" and \"always analyse as primary\" is important and sometimes lost in advocacy for outcome standardisation. More broadly, the choice of outcome is itself a scientific hypothesis about how an intervention works, and should be justified as such in trial protocols. Simulation studies tailored to the specific disease and intervention context can help quantify the consequences of outcome choice and should become a routine part of trial design.": (
            "We do not think this complexity should discourage core outcome sets for SAB. Standardised outcomes remain essential for comparability and should be reported in all trials. However, our data suggest they should function as a minimum reporting standard rather than as a constraint on the primary endpoint. Trialists should be encouraged to select a primary outcome matched to the hypothesised mechanism while still reporting the core set for cross-trial comparison. This distinction between \"always report\" and \"always analyse as primary\" is important and is sometimes lost in advocacy for outcome standardisation. More broadly, the choice of outcome is itself a scientific hypothesis about how an intervention works and should be justified as such in trial protocols. Simulation studies tailored to the specific disease and intervention context can help quantify the consequences of outcome choice and should become a routine part of trial design."
        ),
        "The principal limitation of any simulation study is that data-generating mechanisms may not reflect clinical reality. Whether binary analyses outperform ordinal analyses depends entirely on how the data is simulated, and we urge caution in interpreting our results beyond the demonstration of magnitude of sample size requirements and the importance of outcome design. Our primary parameters derive from Swets et al., whose latent class analyses span multiple cohorts from the UK, Spain, Netherlands, and USA. A recent independent study identifying a largely similar set of subphenotypes provides some reassurance about their validity, but the subphenotype model assigns all patients to a class - in reality some will be unclassifiable, and future work should develop improved assignment models. Our misclassification model assumed random redistribution proportional to subphenotype frequency, which is probably the best case; structured misclassification between clinically similar subphenotypes (for example, between the low-mortality subphenotypes B and E) could produce worse and less predictable bias patterns. The 50% winner's curse shrinkage applied to effect estimates is defensible but arbitrary - if true effects are even smaller, the picture becomes bleaker still.": (
            "The principal limitation of any simulation study is that its data-generating mechanisms may not reflect clinical reality. Whether binary analyses outperform ordinal analyses depends strongly on how outcomes are generated, and our results should therefore be interpreted primarily as illustrating the scale of the sample-size problem and the importance of outcome design. Our primary parameters derive from Swets et al., whose latent class analyses span multiple cohorts from the UK, Spain, the Netherlands, and the USA. A recent independent study identifying a broadly similar set of subphenotypes provides some reassurance about their validity, but the current subphenotype model still assigns every patient to a class; in reality, some patients will be difficult to classify, and future work should develop improved assignment models. Our misclassification model also assumed random redistribution proportional to subphenotype frequency, which is probably optimistic. Structured misclassification between clinically similar subphenotypes, for example between the low-mortality subphenotypes B and E, could produce worse and less predictable bias patterns. Finally, the 50% winner's-curse shrinkage applied to the subgroup effect estimates is defensible but arbitrary; if the true effects are smaller, the conclusions become correspondingly less optimistic."
        ),
        "Our results suggest a three-pronged strategy for advancing stratified trials in SAB. First, diagnostic classifiers must be improved towards near-perfect performance (sensitivity and specificity exceeding 95%) to make even moderately rare subphenotypes viable targets. Second, trial resources should pragmatically focus on subphenotypes with favourable prevalence and event-rate characteristics, as attempting to power trials for rare, low-event subphenotypes with current technology is likely to fail. Third, and perhaps most importantly, outcome measurement must be optimised and matched to the plausible mechanism of the intervention under study. No single outcome will be optimal across all interventions and subphenotypes, and accepting this will require a shift in how infection trials are designed, reported, and interpreted.": (
            "Our results suggest a three-pronged strategy for advancing stratified trials in SAB. First, diagnostic classifiers must improve towards near-perfect performance if even moderately rare subphenotypes are to become viable trial targets. Second, trial resources should pragmatically focus on subphenotypes with favourable prevalence and event-rate characteristics, because attempts to power trials for rare, low-event subphenotypes with current methods are likely to fail. Third, and perhaps most importantly, outcome measurement should be matched to the plausible mechanism of the intervention under study. No single outcome will be optimal across all interventions and subphenotypes, and accepting this will require a shift in how infection trials are designed, reported, and interpreted."
        ),
    }

    doc = Document(str(DOC))
    changed = 0
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt in replacements:
            p.text = replacements[txt]
            changed += 1

    apply_arial(doc)
    doc.save(str(DOC))

    print(f"Updated paragraphs: {changed}")
    print(f"Applied font: Arial")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
