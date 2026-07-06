#!/usr/bin/env python3
"""Sync key manuscript placeholders with current simulation outputs."""

import csv
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
MAIN_METRICS = ROOT / "results" / "main" / "data" / "main_binary_metrics.tsv"
TABLE1_PATH = ROOT / "results" / "paper" / "final" / "tables" / "Table1_main_parameters.tsv"
TABLE2_MAIN_PATH = ROOT / "results" / "paper" / "final" / "tables" / "Table2_main_power_grid.tsv"
TABLES8_PATH = ROOT / "results" / "paper" / "final" / "tables" / "TableS8_ordinal_PO_vs_nonPO_power_bias.tsv"


def resolve_doc_path(root: Path) -> Path:
    paper_dir = root / "paper"
    env_doc = __import__("os").environ.get("MANUSCRIPT_DOC")
    if env_doc:
        return Path(env_doc)

    preferred = paper_dir / "SAB_HET_2.1_JU.docx"
    if preferred.exists():
        return preferred

    raise FileNotFoundError(f"No manuscript doc found in {paper_dir}")


DOC_PATH = resolve_doc_path(ROOT)


def read_tsv(path: Path):
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f, delimiter="\t"))


def fmt3(x: float) -> str:
    return f"{x:.3f}"


def insert_paragraph_before(target_paragraph, text: str):
    new_p = target_paragraph._parent.add_paragraph(text)
    target_paragraph._p.addprevious(new_p._p)
    return new_p


def insert_table_after_paragraph(doc: Document, paragraph, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    tbl = table._element
    paragraph._element.addnext(tbl)
    return table


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def remove_legacy_ordinal_table(doc: Document) -> None:
    for p in list(doc.paragraphs):
        txt = (p.text or "").strip()
        if txt.startswith("Table 2: Ordinal PO vs non-PO performance"):
            remove_paragraph(p)
            break
    if len(doc.tables) > 2:
        doc.tables[2]._element.getparent().remove(doc.tables[2]._element)


def main() -> None:
    main_rows = read_tsv(MAIN_METRICS)
    table1_rows = read_tsv(TABLE1_PATH)
    table2_rows = read_tsv(TABLE2_MAIN_PATH)
    tableS8_rows = read_tsv(TABLES8_PATH)

    power_lookup = {}
    for row in main_rows:
        if float(row["alpha"]) != 0.05 or float(row["accuracy"]) != 1.0:
            continue
        grp = row["group"]
        n = int(float(row["sample_size"]))
        p = row["power"]
        if grp != "A" and p not in ("", "NA", None):
            power_lookup[(n, grp)] = float(p)

    getp = lambda n, g: power_lookup[(n, g)]
    p_b_500 = getp(500, "B")
    p_b_5000 = getp(5000, "B")
    p_b_10000 = getp(10000, "B")
    p_b_20000 = getp(20000, "B")
    p_c_20000 = getp(20000, "C")
    p_d_20000 = getp(20000, "D")
    p_e_20000 = getp(20000, "E")

    type1_a_20000 = [
        float(r["type1"])
        for r in main_rows
        if int(float(r["sample_size"])) == 20000
        and r["group"] == "A"
        and float(r["alpha"]) == 0.05
        and r["type1"] not in ("", "NA", None)
    ]
    type1_min = min(type1_a_20000)
    type1_max = max(type1_a_20000)

    def pick_s8(dgm: str, group: str, model: str, acc: str = "100%"):
        return next(
            r
            for r in tableS8_rows
            if r["Classification accuracy"] == acc
            and r["Data-generating mechanism"] == dgm
            and r["Subphenotype"] == group
            and r["Model"] == model
        )

    po_label = "Proportional-odds"
    np_label = "Death-only non-proportional"
    grp = lambda x: x

    po_b_bin = float(pick_s8(po_label, grp("B"), "Binary (death)")["Power"])
    po_b_ord = float(pick_s8(po_label, grp("B"), "Ordinal (polr)")["Power"])
    po_c_bin = float(pick_s8(po_label, grp("C"), "Binary (death)")["Power"])
    po_c_ord = float(pick_s8(po_label, grp("C"), "Ordinal (polr)")["Power"])
    po_d_bin = float(pick_s8(po_label, grp("D"), "Binary (death)")["Power"])
    po_d_ord = float(pick_s8(po_label, grp("D"), "Ordinal (polr)")["Power"])
    po_e_bin = float(pick_s8(po_label, grp("E"), "Binary (death)")["Power"])
    po_e_ord = float(pick_s8(po_label, grp("E"), "Ordinal (polr)")["Power"])

    np_b_bin = float(pick_s8(np_label, grp("B"), "Binary (death)")["Power"])
    np_b_ord = float(pick_s8(np_label, grp("B"), "Ordinal (polr)")["Power"])
    np_c_bin = float(pick_s8(np_label, grp("C"), "Binary (death)")["Power"])
    np_c_ord = float(pick_s8(np_label, grp("C"), "Ordinal (polr)")["Power"])
    np_d_bin = float(pick_s8(np_label, grp("D"), "Binary (death)")["Power"])
    np_d_ord = float(pick_s8(np_label, grp("D"), "Ordinal (polr)")["Power"])
    np_e_bin = float(pick_s8(np_label, grp("E"), "Binary (death)")["Power"])
    np_e_ord = float(pick_s8(np_label, grp("E"), "Ordinal (polr)")["Power"])

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC_PATH.with_name(f"{DOC_PATH.stem}.backup_{stamp}.docx")
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))
    remove_legacy_ordinal_table(doc)

    add_data = (
        f"At 100% classification accuracy, power for subgroup B rose from {fmt3(p_b_500)} (n=500) "
        f"to {fmt3(p_b_5000)} (n=5,000), {fmt3(p_b_10000)} (n=10,000), and {fmt3(p_b_20000)} (n=20,000). "
        f"At n=20,000, power remained below 50% for C ({fmt3(p_c_20000)}), "
        f"D ({fmt3(p_d_20000)}), and E ({fmt3(p_e_20000)})."
    )

    add_type1 = (
        f"Type I error in the null subgroup (A) remained close to nominal in the main simulation "
        f"(n=20,000), ranging from {100 * type1_min:.1f}% at 100% accuracy to {100 * type1_max:.1f}% at 70% accuracy."
    )

    add_ord = (
        f"At n=3,000 and 100% accuracy, under proportional odds the ordinal model was more powerful than binary in all non-null subgroups: "
        f"B ({fmt3(po_b_ord)} vs {fmt3(po_b_bin)}), C ({fmt3(po_c_ord)} vs {fmt3(po_c_bin)}), "
        f"D ({fmt3(po_d_ord)} vs {fmt3(po_d_bin)}), and E ({fmt3(po_e_ord)} vs {fmt3(po_e_bin)}). "
        f"Under non-proportional death-only effects, binary death analysis was more powerful for B ({fmt3(np_b_bin)} vs {fmt3(np_b_ord)}), "
        f"C ({fmt3(np_c_bin)} vs {fmt3(np_c_ord)}), D ({fmt3(np_d_bin)} vs {fmt3(np_d_ord)}), "
        f"and E ({fmt3(np_e_bin)} vs {fmt3(np_e_ord)})."
    )

    ord_discussion = (
        "When the ordinal outcome was appropriate, the ordinal model substantially increased power while maintaining Type I error control "
        "(Figure 5; Supplementary Table S8). Under misclassification, ordinal estimates also tended to show lower absolute bias than binary estimates, "
        "because information was distributed across multiple cumulative thresholds rather than being concentrated entirely at the death threshold."
    )

    for p in doc.paragraphs:
        txt = p.text or ""
        stripped = txt.strip()

        if stripped.startswith("Staphylococcus aureus bacteraemia (SAB) is clinically heterogeneous. Recently identified, reproducible subphenotypes suggest"):
            p.text = (
                "Staphylococcus aureus bacteraemia (SAB) is clinically heterogeneous. Recently identified, reproducible "
                "subphenotypes, that is, recurrent clinical patient groups with distinct characteristics and outcomes, "
                "suggest potential for heterogeneous treatment effects (HTE), but the impact of patient misclassification "
                "on detecting these effects is unclear, and strategies to improve detection of HTE are not yet identified."
            )
        if stripped.startswith("The statistical power to detect subphenotype-specific treatment effects in a standard two-arm RCT"):
            p.text = (
                "The statistical power to detect subphenotype-specific treatment effects in a standard two-arm RCT, "
                "assuming perfect classification and the conservatively shrunk effects used here, was highly dependent "
                "on subphenotype prevalence, baseline event rate, and effect size. The results, summarised in Table 2 "
                "and Figure 1, show that subphenotype B was readily detectable at realistic sample sizes, reaching "
                "essentially complete power by n=5,000. In contrast, C and D remained only modestly powered even at "
                "larger sample sizes, and E remained difficult under binary mortality analyses despite becoming clearly "
                "more plausible than in earlier iterations of the simulation. Fuller operating characteristics are "
                "provided in Supplementary Table S1. As expected in a traditional two-arm RCT, Type I error was well "
                "calibrated and bias was minimal."
            )
        if "Table 3 and Figure 1" in txt:
            p.text = txt.replace("Table 3 and Figure 1", "Table 2 and Figure 1")
        if stripped.startswith("Table 3: Statistical power to detect"):
            p.text = (
                "Table 2: Statistical power to detect a subphenotype-specific effect at each sample size n. "
                "Power was estimated from 2,000 simulated replicates. Subphenotype A has no treatment effect, "
                "so power is not the appropriate metric."
            )
        if stripped == "ADD TYPE1 ERROR":
            p.text = add_type1
        if stripped.startswith("Ordinal outcomes rescue power and reduce bias when correctly specified"):
            p.text = "Ordinal outcomes can substantially improve power, but the tradeoff depends on the outcome scale and treatment-effect mechanism"
        if stripped.startswith("When the ordinal outcome was appropriate"):
            p.text = ord_discussion
        if stripped.startswith("At n=3,000 and 100% accuracy, under proportional odds"):
            p.text = add_ord
        if stripped.startswith("These results show the strengths and tradeoffs of ordinal outcomes."):
            p.text = (
                "These results show the strengths and tradeoffs of ordinal outcomes. When the ordinal structure "
                "reasonably reflects the treatment effect, they can be much more powerful and less biased than "
                "binary mortality analyses. When the effect is concentrated mainly at death, binary analysis remains "
                "preferable, but the degree of ordinal underperformance depends on how the ordinal scale is "
                "constructed. In particular, our 6-point main analysis performed slightly better under "
                "proportional-odds scenarios, whereas the 5-point sensitivity analysis was somewhat more robust "
                "under death-only non-proportional scenarios. This tradeoff is shown directly in Supplementary "
                "Figure S1, with additional ordinal performance and calibration detail in Supplementary Tables S4-S8."
            )
        if stripped.startswith("Figure 4:"):
            p.text = (
                "Figure 4: Ordinal outcome category shifts under proportional-odds (A) and non-proportional "
                "(death-only) effects (B) for subgroup B. Stacked bars show the control and treatment "
                "distributions across the six ordered categories (death, ICU/ventilated, still hospitalised, discharged to rehab, "
                "discharged with complications, discharged well). Within the non-proportional panel, survivor categories retain their "
                "relative composition and change only through rescaling after the treatment-associated shift in death."
            )
        if stripped.startswith("It is important to understand how conditional these estimates are on the cohort under study"):
            p.text = (
                "It is important to understand how conditional these estimates are on the cohort under study, because "
                "sample-size calculations are made before a trial starts and the extent to which subgroup prevalence "
                "and mortality vary across settings is uncertain. We therefore reran the analyses using data from several "
                "cohorts with differing subgroup prevalences and mortality, assuming perfect classification. For most "
                "subgroups, the required sample size varied substantially by setting, in some cases by as much as "
                "five-fold (Figure 2; Supplementary Table S2)."
            )
        if stripped.startswith("We evaluated the impact of five hypothetical diagnostic tests ranging from mediocre to near-perfect accuracy"):
            p.text = (
                "We evaluated the impact of five hypothetical diagnostic tests ranging from mediocre to near-perfect "
                "accuracy on the number needed to screen (NNS) and number needed to randomise (NNR). For subgroup B, "
                "which combines a large treatment effect (OR 4.3 after shrinkage) with moderate frequency (13%), enrichment "
                "remained realistic: NNS ranged from about 1,100 with a near-perfect test (99%/99%) to about 3,300 with "
                "a balanced test (80%/80%), with corresponding NNR values of roughly 140 to 930 participants (Figure 3; "
                "Supplementary Table S3). Biased estimates were still present, as expected. In contrast, for the other "
                "subgroups enrichment remained unattractive or infeasible even with excellent classifiers. Type I "
                "error in the null subgroup (A) was 5.0% at 100% accuracy and rose to 7.0% at 70% accuracy in the "
                "deliberately large n=20,000 simulation."
            )
        if "[ADD DATA]" in txt:
            p.text = txt.replace("[ADD DATA]", add_data)

    if not any((p.text or "").strip().startswith("Figure 5:") for p in doc.paragraphs):
        anchor = next((p for p in doc.paragraphs if (p.text or "").strip().startswith("At n=3,000 and 100% accuracy")), None)
        if anchor is not None:
            insert_paragraph_before(
                anchor,
                "Figure 5: Binary versus ordinal performance under proportional-odds and non-proportional "
                "(death-only) ordinal data-generating mechanisms, stratified by classification accuracy. "
                "Panels A and C show replicate scaled error at n=20,000; panels B and D show power comparisons at n=3,000."
            )

    t1 = doc.tables[0]
    t1.cell(0, 0).text = "Subgroup"
    t1.cell(0, 1).text = "Frequency in ARREST placebo arm (%)"
    t1.cell(0, 2).text = "ARREST 84-day mortality (%)"
    t1.cell(0, 3).text = "ARREST OR for 84-day mortality"
    t1.cell(0, 4).text = "Shrunk OR for 84-day mortality"
    for ridx, row in enumerate(table1_rows, start=1):
        if ridx >= len(t1.rows):
            break
        t1.cell(ridx, 0).text = row["Subphenotype"]
        t1.cell(ridx, 1).text = row["Frequency in ARREST placebo arm (%)"]
        t1.cell(ridx, 2).text = row["Baseline 84-day mortality across both arms (%)"]
        t1.cell(ridx, 3).text = row["Original OR for 84-day mortality"]
        t1.cell(ridx, 4).text = row["Conservative OR for 84-day mortality"]

    table2_caption = next(
        p for p in doc.paragraphs
        if (p.text or "").strip().startswith("Table 2: Statistical power to detect")
    )
    if len(doc.tables) > 1:
        doc.tables[1]._element.getparent().remove(doc.tables[1]._element)
    table2_doc_rows = [
        ["Total trial size", "Subphenotype A (null)", "Subphenotype B", "Subphenotype C", "Subphenotype D", "Subphenotype E"],
        ["", "A", "B", "C", "D", "E"],
    ]
    for row in table2_rows:
        table2_doc_rows.append([
            row["Total trial size"], row["A"], row["B"], row["C"], row["D"], row["E"]
        ])
    insert_table_after_paragraph(doc, table2_caption, table2_doc_rows)

    doc.save(str(DOC_PATH))
    print(f"Updated: {DOC_PATH}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
